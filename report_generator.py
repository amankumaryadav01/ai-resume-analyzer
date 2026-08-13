"""
report_generator.py

Builds downloadable PDF and DOCX versions of:
1. ATS Diagnostic Report (analysis result)
2. Improved Resume (rewritten ATS-optimized resume)
"""

import io
from datetime import datetime

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.graphics.shapes import Drawing, Rect

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_NAME = "AI Resume Copilot"
REPORT_TITLE = "ATS Diagnostic Report"


def _band_color_hex(value: int) -> str:
    """0-40 red / 40-70 amber / 70-100 green."""
    if value >= 70:
        return "#1FAE6B"
    if value >= 40:
        return "#D97706"
    return "#DC2626"


# =====================================================================
# PDF HELPERS & ATS REPORT GENERATOR
# =====================================================================

def _register_fonts():
    return "Helvetica", "Helvetica-Bold"


def _build_styles(body_font, bold_font):
    return {
        "BrandTitle": ParagraphStyle(
            "BrandTitle", fontName=bold_font, fontSize=26,
            textColor=colors.HexColor("#14161B"), alignment=TA_CENTER, spaceAfter=6,
        ),
        "Subtitle": ParagraphStyle(
            "Subtitle", fontName=body_font, fontSize=13,
            textColor=colors.HexColor("#6366F1"), alignment=TA_CENTER,
        ),
        "Section": ParagraphStyle(
            "Section", fontName=bold_font, fontSize=14,
            textColor=colors.HexColor("#14161B"), spaceBefore=6, spaceAfter=8,
        ),
        "Body": ParagraphStyle(
            "Body", fontName=body_font, fontSize=10.5, leading=15,
            textColor=colors.HexColor("#14161B"),
        ),
        "Muted": ParagraphStyle(
            "Muted", fontName=body_font, fontSize=10.5,
            textColor=colors.HexColor("#6B7280"),
        ),
    }


def _score_bar(value: int, width=380, height=16):
    color_hex = _band_color_hex(value)
    d = Drawing(width, height)
    d.add(Rect(0, 0, width, height, fillColor=colors.HexColor("#EDEEF3"), strokeColor=None, rx=6, ry=6))
    fill_w = max(6, width * max(0, min(100, value)) / 100)
    d.add(Rect(0, 0, fill_w, height, fillColor=colors.HexColor(color_hex), strokeColor=None, rx=6, ry=6))
    return d


def _bullets(items, styles, empty_text):
    if not items:
        return [Paragraph(empty_text, styles["Muted"])]
    return [Paragraph(f"•&nbsp;&nbsp;{text}", styles["Body"]) for text in items]


class _NumberedCanvas(pdfcanvas.Canvas):
    def __init__(self, *args, **kwargs):
        pdfcanvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer(num_pages)
            pdfcanvas.Canvas.showPage(self)
        pdfcanvas.Canvas.save(self)

    def _draw_footer(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#9AA0AC"))
        self.drawString(0.9 * inch, 0.55 * inch, f"{BRAND_NAME} — {REPORT_TITLE}")
        self.drawRightString(LETTER[0] - 0.9 * inch, 0.55 * inch, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()


def generate_pdf_bytes(result: dict, resume_filename: str) -> bytes:
    """Builds ATS Analysis Report PDF."""
    body_font, bold_font = _register_fonts()
    styles = _build_styles(body_font, bold_font)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        title=REPORT_TITLE,
    )

    story = []

    # Title page
    story.append(Spacer(1, 1.3 * inch))
    story.append(Paragraph(BRAND_NAME, styles["BrandTitle"]))
    story.append(Paragraph(REPORT_TITLE, styles["Subtitle"]))
    story.append(Spacer(1, 0.5 * inch))
    story.append(HRFlowable(width="100%", color=colors.HexColor("#E1E4EA"), thickness=1))
    story.append(Spacer(1, 0.35 * inch))

    name = result.get("name") or "Candidate"
    email = result.get("email") or "Not provided"
    generated = datetime.now().strftime("%B %d, %Y  •  %I:%M %p")

    meta_table = Table(
        [
            ["Candidate", name],
            ["Email", email],
            ["Resume File", resume_filename or "N/A"],
            ["Generated", generated],
        ],
        colWidths=[1.6 * inch, 4.0 * inch],
    )
    meta_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), bold_font),
        ("FONTNAME", (1, 0), (1, -1), body_font),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#6B7280")),
        ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#14161B")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(meta_table)
    story.append(PageBreak())

    # Overview
    ats = result.get("ats_score", 0)
    match = result.get("match_percentage", 0)
    matching = result.get("matching_skills") or []
    missing = result.get("missing_skills") or []

    story.append(Paragraph("Overview", styles["Section"]))
    story.append(Paragraph(f"ATS Score — {ats}/100", styles["Body"]))
    story.append(Spacer(1, 4))
    story.append(_score_bar(ats))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Match Percentage — {match}%", styles["Body"]))
    story.append(Spacer(1, 4))
    story.append(_score_bar(match))
    story.append(Spacer(1, 16))

    kpi_table = Table(
        [
            ["ATS Score", "Match %", "Skills Matched", "Skills Missing"],
            [str(ats), f"{match}%", str(len(matching)), str(len(missing))],
        ],
        colWidths=[1.2 * inch] * 4,
    )
    kpi_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), bold_font),
        ("FONTNAME", (0, 1), (-1, 1), body_font),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F2FA")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E1E4EA")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 20))

    story.append(Paragraph("AI Summary", styles["Section"]))
    story.append(Paragraph(result.get("summary") or "No summary generated.", styles["Body"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Matching Skills", styles["Section"]))
    story.extend(_bullets(matching, styles, "No matching skills identified."))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Missing Skills", styles["Section"]))
    story.extend(_bullets(missing, styles, "No missing skills."))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Strengths", styles["Section"]))
    story.extend(_bullets(result.get("strengths") or [], styles, "None listed."))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Weaknesses", styles["Section"]))
    story.extend(_bullets(result.get("weaknesses") or [], styles, "None listed."))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Suggestions", styles["Section"]))
    suggestions = result.get("suggestions") or []
    if suggestions:
        for i, s in enumerate(suggestions, start=1):
            story.append(Paragraph(f"{i}. {s}", styles["Body"]))
            story.append(Spacer(1, 4))

    doc.build(story, canvasmaker=_NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()


# =====================================================================
# DOCX ATS REPORT GENERATOR
# =====================================================================

def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_score_bar(doc, label: str, value: int, segments: int = 20):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} — {value}/100")
    run.bold = True
    filled = round(segments * max(0, min(100, value)) / 100)
    fill_color = _band_color_hex(value).lstrip("#")
    table = doc.add_table(rows=1, cols=segments)
    table.autofit = False
    for i in range(segments):
        cell = table.cell(0, i)
        cell.width = Cm(0.32)
        _set_cell_shading(cell, fill_color if i < filled else "E5E7F0")
    doc.add_paragraph()


def generate_docx_bytes(result: dict, resume_filename: str) -> bytes:
    """Builds ATS Analysis Report DOCX."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    header_p = section.header.paragraphs[0]
    header_run = header_p.add_run(f"{BRAND_NAME} — {REPORT_TITLE}")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x67, 0x6E, 0x7C)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("Page ").font.size = Pt(9)
    _add_page_number_field(footer_p)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(BRAND_NAME)
    title_run.font.size = Pt(28)
    title_run.font.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(REPORT_TITLE)
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

    doc.add_paragraph()

    name = result.get("name") or "Candidate"
    email = result.get("email") or "Not provided"
    generated = datetime.now().strftime("%B %d, %Y  •  %I:%M %p")

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Light List Accent 1"
    meta_rows = [
        ("Candidate", name),
        ("Email", email),
        ("Resume File", resume_filename or "N/A"),
        ("Generated", generated),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = str(value)

    doc.add_page_break()

    ats = result.get("ats_score", 0)
    match = result.get("match_percentage", 0)
    matching = result.get("matching_skills") or []
    missing = result.get("missing_skills") or []

    doc.add_heading("Overview", level=1)
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Light Grid Accent 1"
    headers = ["ATS Score", "Match %", "Skills Matched", "Skills Missing"]
    values = [str(ats), f"{match}%", str(len(matching)), str(len(missing))]
    for i, h in enumerate(headers):
        kpi_table.cell(0, i).text = h
    for i, v in enumerate(values):
        kpi_table.cell(1, i).text = v

    doc.add_paragraph()
    _add_score_bar(doc, "ATS Score", ats)
    _add_score_bar(doc, "Match Percentage", match)

    doc.add_heading("AI Summary", level=1)
    doc.add_paragraph(result.get("summary") or "No summary generated.")

    doc.add_heading("Matching Skills", level=1)
    if matching:
        for m in matching:
            doc.add_paragraph(m, style="List Bullet")
    else:
        doc.add_paragraph("No matching skills identified.")

    doc.add_heading("Missing Skills", level=1)
    if missing:
        for ms in missing:
            doc.add_paragraph(ms, style="List Bullet")
    else:
        doc.add_paragraph("No missing skills.")

    doc.add_heading("Strengths", level=1)
    for st in result.get("strengths") or []:
        doc.add_paragraph(st, style="List Bullet")

    doc.add_heading("Weaknesses", level=1)
    for wk in result.get("weaknesses") or []:
        doc.add_paragraph(wk, style="List Bullet")

    doc.add_heading("Suggestions", level=1)
    for i, s in enumerate(result.get("suggestions") or [], start=1):
        doc.add_paragraph(f"{i}. {s}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =====================================================================
# IMPROVED RESUME EXPORTERS (PDF & DOCX)
# =====================================================================

def generate_resume_pdf(data: dict) -> bytes:
    """Generates a professional PDF version of the rewritten resume."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=0.65 * inch, bottomMargin=0.65 * inch,
        leftMargin=0.65 * inch, rightMargin=0.65 * inch
    )

    PRIMARY = colors.HexColor("#1A365D")
    TEXT = colors.HexColor("#2D3748")

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("RName", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=PRIMARY, alignment=1)
    contact_style = ParagraphStyle("RContact", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12, textColor=TEXT, alignment=1, spaceAfter=6)
    section_style = ParagraphStyle("RSec", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=PRIMARY, spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle("RBody", parent=styles["Normal"], fontName="Helvetica", fontSize=9.5, leading=13.5, textColor=TEXT, spaceAfter=3)
    bullet_style = ParagraphStyle("RBullet", parent=body_style, leftIndent=12, firstLineIndent=-8, spaceAfter=2)

    story = []

    # Name & Contact
    story.append(Paragraph(data.get("name", "Applicant Name"), name_style))
    contact = data.get("contact", {})
    c_parts = [contact.get(k) for k in ["email", "phone", "location", "linkedin", "portfolio"] if contact.get(k)]
    story.append(Paragraph(" | ".join(c_parts), contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceAfter=6))

    def add_sec(title):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CBD5E0"), spaceAfter=4))

    # Summary
    if data.get("summary"):
        add_sec("Professional Summary")
        story.append(Paragraph(data["summary"], body_style))

    # Skills
    if data.get("skills"):
        add_sec("Skills")
        for sk in data["skills"]:
            story.append(Paragraph(f"<b>{sk.get('category', 'Category')}:</b> {', '.join(sk.get('skills', []))}", body_style))

    # Experience
    if data.get("experience"):
        add_sec("Experience")
        for exp in data["experience"]:
            header = f"<b>{exp.get('title')}</b> — {exp.get('company')} ({exp.get('dates', '')})"
            story.append(Paragraph(header, body_style))
            for b in exp.get("bullets", []):
                story.append(Paragraph(f"• {b}", bullet_style))
            story.append(Spacer(1, 3))

    # Projects
    if data.get("projects"):
        add_sec("Projects")
        for proj in data["projects"]:
            header = f"<b>{proj.get('name')}</b> | <i>{proj.get('technologies', '')}</i>"
            story.append(Paragraph(header, body_style))
            for b in proj.get("bullets", []):
                story.append(Paragraph(f"• {b}", bullet_style))
            story.append(Spacer(1, 3))

    # Education
    if data.get("education"):
        add_sec("Education")
        for edu in data["education"]:
            story.append(Paragraph(f"<b>{edu.get('degree')}</b> — {edu.get('institution')} ({edu.get('dates', '')})", body_style))

    # Certifications
    if data.get("certifications"):
        add_sec("Certifications")
        for cert in data["certifications"]:
            story.append(Paragraph(f"• {cert}", bullet_style))

    doc.build(story, canvasmaker=_NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()


def generate_resume_docx(data: dict) -> bytes:
    """Generates a professional DOCX version of the rewritten resume."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    COLOR_PRIMARY = RGBColor(26, 54, 93)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(data.get("name", "Applicant Name"))
    r_name.font.size = Pt(18)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY

    contact = data.get("contact", {})
    c_parts = [contact.get(k) for k in ["email", "phone", "location", "linkedin", "portfolio"] if contact.get(k)]
    if c_parts:
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c.add_run(" | ".join(c_parts))

    def add_sec_title(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title.upper())
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY

    if data.get("summary"):
        add_sec_title("Professional Summary")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        add_sec_title("Skills")
        for sk in data["skills"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{sk.get('category', 'Category')}: ")
            r.bold = True
            p.add_run(", ".join(sk.get("skills", [])))

    if data.get("experience"):
        add_sec_title("Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{exp.get('title')} — {exp.get('company')} ({exp.get('dates', '')})")
            r.bold = True
            for b in exp.get("bullets", []):
                doc.add_paragraph(b, style='List Bullet')

    if data.get("projects"):
        add_sec_title("Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{proj.get('name')} | {proj.get('technologies', '')}")
            r.bold = True
            for b in proj.get("bullets", []):
                doc.add_paragraph(b, style='List Bullet')

    if data.get("education"):
        add_sec_title("Education")
        for edu in data["education"]:
            doc.add_paragraph(f"{edu.get('degree')} — {edu.get('institution')} ({edu.get('dates', '')})")

    if data.get("certifications"):
        add_sec_title("Certifications")
        for cert in data["certifications"]:
            doc.add_paragraph(cert, style='List Bullet')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()