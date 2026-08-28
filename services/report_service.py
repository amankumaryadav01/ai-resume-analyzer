"""
report_service.py

Generates:
1. ATS diagnostic PDF
2. ATS diagnostic DOCX
3. Optimized Resume PDF
4. Optimized Resume DOCX

The optimized resume PDF/DOCX uses the SAME optimized dictionary
that is displayed in Streamlit. No second AI generation happens here.
"""

import io
import functools
from datetime import datetime

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    HRFlowable,
)
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.graphics.shapes import Drawing, Rect

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_NAME = "AI Resume Copilot"
REPORT_TITLE = "ATS Diagnostic Report"


# ============================================================
# COMMON HELPERS
# ============================================================

def _band_color_hex(value: int) -> str:
    if value >= 70:
        return "#1FAE6B"
    if value >= 40:
        return "#D97706"
    return "#DC2626"


def _register_fonts():
    return "Helvetica", "Helvetica-Bold"


def _build_styles(body_font, bold_font):
    return {
        "BrandTitle": ParagraphStyle(
            "BrandTitle",
            fontName=bold_font,
            fontSize=26,
            leading=31,
            textColor=colors.HexColor("#14161B"),
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "Subtitle": ParagraphStyle(
            "Subtitle",
            fontName=body_font,
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#6366F1"),
            alignment=TA_CENTER,
        ),
        "Section": ParagraphStyle(
            "Section",
            fontName=bold_font,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#14161B"),
            spaceBefore=6,
            spaceAfter=8,
        ),
        "Body": ParagraphStyle(
            "Body",
            fontName=body_font,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#14161B"),
        ),
        "Muted": ParagraphStyle(
            "Muted",
            fontName=body_font,
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#6B7280"),
        ),
    }


def _score_bar(value: int, width=380, height=16):
    color_hex = _band_color_hex(value)

    d = Drawing(width, height)

    d.add(
        Rect(
            0,
            0,
            width,
            height,
            fillColor=colors.HexColor("#EDEEF3"),
            strokeColor=None,
            rx=6,
            ry=6,
        )
    )

    fill_w = max(
        6,
        width * max(0, min(100, value)) / 100,
    )

    d.add(
        Rect(
            0,
            0,
            fill_w,
            height,
            fillColor=colors.HexColor(color_hex),
            strokeColor=None,
            rx=6,
            ry=6,
        )
    )

    return d


def _bullets(items, styles, empty_text=""):
    if not items:
        if empty_text:
            return [Paragraph(empty_text, styles["Muted"])]
        return []

    return [
        Paragraph(
            f"•&nbsp;&nbsp;{str(text)}",
            styles["Body"],
        )
        for text in items
        if str(text).strip()
    ]


# ============================================================
# PAGE NUMBER CANVAS
# ============================================================

class _NumberedCanvas(pdfcanvas.Canvas):

    def __init__(self, *args, footer_label=None, **kwargs):
        self._footer_label = (
            footer_label
            or f"{BRAND_NAME} — {REPORT_TITLE}"
        )

        pdfcanvas.Canvas.__init__(
            self,
            *args,
            **kwargs,
        )

        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(
            dict(self.__dict__)
        )
        self._startPage()

    def save(self):
        page_count = len(
            self._saved_page_states
        )

        for state in self._saved_page_states:
            self.__dict__.update(state)

            self._draw_footer(page_count)

            pdfcanvas.Canvas.showPage(self)

        pdfcanvas.Canvas.save(self)

    def _draw_footer(self, page_count):
        self.saveState()

        self.setFont(
            "Helvetica",
            8,
        )

        self.setFillColor(
            colors.HexColor("#9AA0AC")
        )

        self.drawString(
            0.9 * inch,
            0.55 * inch,
            self._footer_label,
        )

        self.drawRightString(
            LETTER[0] - 0.9 * inch,
            0.55 * inch,
            f"Page {self._pageNumber} of {page_count}",
        )

        self.restoreState()


# ============================================================
# ATS DIAGNOSTIC PDF
# ============================================================

def generate_pdf_bytes(
    result: dict,
    resume_filename: str,
) -> bytes:

    body_font, bold_font = _register_fonts()
    styles = _build_styles(
        body_font,
        bold_font,
    )

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        title=REPORT_TITLE,
    )

    story = []

    # --------------------------------------------------------
    # TITLE
    # --------------------------------------------------------

    story.append(
        Spacer(1, 1.3 * inch)
    )

    story.append(
        Paragraph(
            BRAND_NAME,
            styles["BrandTitle"],
        )
    )

    story.append(
        Paragraph(
            REPORT_TITLE,
            styles["Subtitle"],
        )
    )

    story.append(
        Spacer(1, 0.5 * inch)
    )

    story.append(
        HRFlowable(
            width="100%",
            color=colors.HexColor("#E1E4EA"),
            thickness=1,
        )
    )

    story.append(
        Spacer(1, 0.35 * inch)
    )

    name = result.get(
        "name",
        "Candidate",
    )

    email = result.get(
        "email",
        "Not provided",
    )

    generated = datetime.now().strftime(
        "%B %d, %Y  •  %I:%M %p"
    )

    meta_table = Table(
        [
            ["Candidate", name],
            ["Email", email],
            [
                "Resume File",
                resume_filename or "N/A",
            ],
            ["Generated", generated],
        ],
        colWidths=[
            1.6 * inch,
            4.0 * inch,
        ],
    )

    meta_table.setStyle(
        TableStyle(
            [
                (
                    "FONTNAME",
                    (0, 0),
                    (0, -1),
                    bold_font,
                ),
                (
                    "FONTNAME",
                    (1, 0),
                    (1, -1),
                    body_font,
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    10,
                ),
                (
                    "TEXTCOLOR",
                    (0, 0),
                    (0, -1),
                    colors.HexColor("#6B7280"),
                ),
                (
                    "TEXTCOLOR",
                    (1, 0),
                    (1, -1),
                    colors.HexColor("#14161B"),
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
            ]
        )
    )

    story.append(meta_table)

    story.append(PageBreak())

    # --------------------------------------------------------
    # OVERVIEW
    # --------------------------------------------------------

    ats = result.get(
        "ats_score",
        0,
    )

    match = result.get(
        "match_percentage",
        0,
    )

    matching = result.get(
        "matching_skills"
    ) or []

    missing = result.get(
        "missing_skills"
    ) or []

    story.append(
        Paragraph(
            "Overview",
            styles["Section"],
        )
    )

    story.append(
        Paragraph(
            f"ATS Score — {ats}/100",
            styles["Body"],
        )
    )

    story.append(
        Spacer(1, 4)
    )

    story.append(
        _score_bar(ats)
    )

    story.append(
        Spacer(1, 12)
    )

    story.append(
        Paragraph(
            f"Match Percentage — {match}%",
            styles["Body"],
        )
    )

    story.append(
        Spacer(1, 4)
    )

    story.append(
        _score_bar(match)
    )

    story.append(
        Spacer(1, 16)
    )

    kpi_table = Table(
        [
            [
                "ATS Score",
                "Match %",
                "Skills Matched",
                "Skills Missing",
            ],
            [
                str(ats),
                f"{match}%",
                str(len(matching)),
                str(len(missing)),
            ],
        ],
        colWidths=[
            1.2 * inch
        ] * 4,
    )

    kpi_table.setStyle(
        TableStyle(
            [
                (
                    "FONTNAME",
                    (0, 0),
                    (-1, 0),
                    bold_font,
                ),
                (
                    "FONTNAME",
                    (0, 1),
                    (-1, 1),
                    body_font,
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    10,
                ),
                (
                    "ALIGN",
                    (0, 0),
                    (-1, -1),
                    "CENTER",
                ),
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.HexColor("#F1F2FA"),
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    colors.HexColor("#E1E4EA"),
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
            ]
        )
    )

    story.append(kpi_table)

    story.append(
        Spacer(1, 20)
    )

    # --------------------------------------------------------
    # AI SUMMARY
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "AI Summary",
            styles["Section"],
        )
    )

    summary = result.get(
        "summary"
    ) or "No summary was generated for this scan."

    story.append(
        Paragraph(
            str(summary),
            styles["Body"],
        )
    )

    story.append(
        Spacer(1, 16)
    )

    # --------------------------------------------------------
    # MATCHING SKILLS
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "Matching Skills",
            styles["Section"],
        )
    )

    story.extend(
        _bullets(
            matching,
            styles,
            "No matching skills identified.",
        )
    )

    story.append(
        Spacer(1, 16)
    )

    # --------------------------------------------------------
    # MISSING SKILLS
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "Missing Skills",
            styles["Section"],
        )
    )

    story.extend(
        _bullets(
            missing,
            styles,
            "No missing skills — full coverage.",
        )
    )

    story.append(
        Spacer(1, 16)
    )

    # --------------------------------------------------------
    # STRENGTHS
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "Strengths",
            styles["Section"],
        )
    )

    story.extend(
        _bullets(
            result.get("strengths") or [],
            styles,
            "None listed.",
        )
    )

    story.append(
        Spacer(1, 16)
    )

    # --------------------------------------------------------
    # WEAKNESSES
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "Weaknesses",
            styles["Section"],
        )
    )

    story.extend(
        _bullets(
            result.get("weaknesses") or [],
            styles,
            "None listed.",
        )
    )

    story.append(
        Spacer(1, 16)
    )

    # --------------------------------------------------------
    # SUGGESTIONS
    # --------------------------------------------------------

    story.append(
        Paragraph(
            "Suggestions",
            styles["Section"],
        )
    )

    suggestions = result.get(
        "suggestions"
    ) or []

    if suggestions:

        for i, suggestion in enumerate(
            suggestions,
            start=1,
        ):
            story.append(
                Paragraph(
                    f"{i}. {suggestion}",
                    styles["Body"],
                )
            )

            story.append(
                Spacer(1, 4)
            )

    else:

        story.append(
            Paragraph(
                "None listed.",
                styles["Body"],
            )
        )

    doc.build(
        story,
        canvasmaker=_NumberedCanvas,
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# DOCX ATS REPORT
# ============================================================

def _add_page_number_field(paragraph):

    run = paragraph.add_run()

    fld_begin = OxmlElement(
        "w:fldChar"
    )

    fld_begin.set(
        qn("w:fldCharType"),
        "begin",
    )

    instr = OxmlElement(
        "w:instrText"
    )

    instr.set(
        qn("xml:space"),
        "preserve",
    )

    instr.text = "PAGE"

    fld_end = OxmlElement(
        "w:fldChar"
    )

    fld_end.set(
        qn("w:fldCharType"),
        "end",
    )

    run._r.append(
        fld_begin
    )

    run._r.append(
        instr
    )

    run._r.append(
        fld_end
    )


def _set_cell_shading(
    cell,
    hex_color: str,
):

    shd = OxmlElement(
        "w:shd"
    )

    shd.set(
        qn("w:val"),
        "clear",
    )

    shd.set(
        qn("w:color"),
        "auto",
    )

    shd.set(
        qn("w:fill"),
        hex_color,
    )

    cell._tc.get_or_add_tcPr().append(
        shd
    )


def _add_score_bar(
    doc,
    label: str,
    value: int,
    segments: int = 20,
):

    p = doc.add_paragraph()

    run = p.add_run(
        f"{label} — {value}/100"
    )

    run.bold = True

    filled = round(
        segments
        * max(0, min(100, value))
        / 100
    )

    fill_color = (
        _band_color_hex(value)
        .lstrip("#")
    )

    table = doc.add_table(
        rows=1,
        cols=segments,
    )

    table.autofit = False

    for i in range(segments):

        cell = table.cell(
            0,
            i,
        )

        cell.width = Cm(0.32)

        _set_cell_shading(
            cell,
            (
                fill_color
                if i < filled
                else "E5E7F0"
            ),
        )

    doc.add_paragraph()


def _add_bullets(
    doc,
    items,
    empty_text,
):

    if not items:

        doc.add_paragraph(
            empty_text
        )

        return

    for item in items:

        doc.add_paragraph(
            str(item),
            style="List Bullet",
        )


def generate_docx_bytes(
    result: dict,
    resume_filename: str,
) -> bytes:

    doc = Document()

    section = doc.sections[0]

    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    header_p = (
        section
        .header
        .paragraphs[0]
    )

    header_run = header_p.add_run(
        f"{BRAND_NAME} — {REPORT_TITLE}"
    )

    header_run.font.size = Pt(9)

    header_run.font.color.rgb = RGBColor(
        0x67,
        0x6E,
        0x7C,
    )

    footer_p = (
        section
        .footer
        .paragraphs[0]
    )

    footer_p.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    footer_p.add_run(
        "Page "
    ).font.size = Pt(9)

    _add_page_number_field(
        footer_p
    )

    title_p = doc.add_paragraph()

    title_p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    title_run = title_p.add_run(
        BRAND_NAME
    )

    title_run.font.size = Pt(28)
    title_run.font.bold = True

    subtitle_p = doc.add_paragraph()

    subtitle_p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle_run = subtitle_p.add_run(
        REPORT_TITLE
    )

    subtitle_run.font.size = Pt(14)

    subtitle_run.font.color.rgb = RGBColor(
        0x63,
        0x66,
        0xF1,
    )

    doc.add_paragraph()

    name = result.get(
        "name",
        "Candidate",
    )

    email = result.get(
        "email",
        "Not provided",
    )

    generated = datetime.now().strftime(
        "%B %d, %Y  •  %I:%M %p"
    )

    meta_table = doc.add_table(
        rows=4,
        cols=2,
    )

    meta_table.style = (
        "Light List Accent 1"
    )

    meta_rows = [
        ("Candidate", name),
        ("Email", email),
        (
            "Resume File",
            resume_filename or "N/A",
        ),
        ("Generated", generated),
    ]

    for i, (
        label,
        value,
    ) in enumerate(meta_rows):

        meta_table.cell(
            i,
            0,
        ).text = label

        meta_table.cell(
            i,
            1,
        ).text = str(value)

    doc.add_page_break()

    ats = result.get(
        "ats_score",
        0,
    )

    match = result.get(
        "match_percentage",
        0,
    )

    matching = result.get(
        "matching_skills"
    ) or []

    missing = result.get(
        "missing_skills"
    ) or []

    doc.add_heading(
        "Overview",
        level=1,
    )

    kpi_table = doc.add_table(
        rows=2,
        cols=4,
    )

    kpi_table.style = (
        "Light Grid Accent 1"
    )

    headers = [
        "ATS Score",
        "Match %",
        "Skills Matched",
        "Skills Missing",
    ]

    values = [
        str(ats),
        f"{match}%",
        str(len(matching)),
        str(len(missing)),
    ]

    for i, header in enumerate(
        headers
    ):
        kpi_table.cell(
            0,
            i,
        ).text = header

    for i, value in enumerate(
        values
    ):
        kpi_table.cell(
            1,
            i,
        ).text = value

    doc.add_paragraph()

    _add_score_bar(
        doc,
        "ATS Score",
        ats,
    )

    _add_score_bar(
        doc,
        "Match Percentage",
        match,
    )

    doc.add_heading(
        "AI Summary",
        level=1,
    )

    doc.add_paragraph(
        result.get("summary")
        or "No summary was generated for this scan."
    )

    doc.add_heading(
        "Matching Skills",
        level=1,
    )

    _add_bullets(
        doc,
        matching,
        "No matching skills identified.",
    )

    doc.add_heading(
        "Missing Skills",
        level=1,
    )

    _add_bullets(
        doc,
        missing,
        "No missing skills — full coverage.",
    )

    doc.add_heading(
        "Strengths",
        level=1,
    )

    _add_bullets(
        doc,
        result.get("strengths") or [],
        "None listed.",
    )

    doc.add_heading(
        "Weaknesses",
        level=1,
    )

    _add_bullets(
        doc,
        result.get("weaknesses") or [],
        "None listed.",
    )

    doc.add_heading(
        "Suggestions",
        level=1,
    )

    suggestions = result.get(
        "suggestions"
    ) or []

    if suggestions:

        for i, suggestion in enumerate(
            suggestions,
            start=1,
        ):

            doc.add_paragraph(
                f"{i}. {suggestion}"
            )

    else:

        doc.add_paragraph(
            "None listed."
        )

    buffer = io.BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# OPTIMIZED RESUME PDF
# ============================================================

def generate_optimized_resume_pdf(
    optimized: dict,
    candidate_name: str = "",
) -> bytes:

    """
    IMPORTANT:

    This function does NOT optimize anything.

    It ONLY renders the already optimized resume.

    The SAME `optimized` dictionary displayed in Streamlit
    should be passed here.

    Therefore:

    Browser Preview == Downloaded PDF
    """

    body_font, bold_font = _register_fonts()

    styles = _build_styles(
        body_font,
        bold_font,
    )

    # --------------------------------------------------------
    # RESUME STYLES
    # --------------------------------------------------------

    heading_style = ParagraphStyle(
        "ResumeHeading",
        fontName=bold_font,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor(
            "#14161B"
        ),
        spaceBefore=14,
        spaceAfter=6,
    )

    name_style = ParagraphStyle(
        "ResumeName",
        fontName=bold_font,
        fontSize=20,
        leading=25,
        textColor=colors.HexColor(
            "#14161B"
        ),
        alignment=TA_CENTER,
        spaceAfter=4,
    )

    role_style = ParagraphStyle(
        "ResumeRole",
        fontName=body_font,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(
            "#6366F1"
        ),
        alignment=TA_CENTER,
        spaceAfter=10,
    )

    entry_title_style = ParagraphStyle(
        "EntryTitle",
        fontName=bold_font,
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor(
            "#14161B"
        ),
        spaceBefore=8,
    )

    entry_meta_style = ParagraphStyle(
        "EntryMeta",
        fontName=body_font,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor(
            "#6B7280"
        ),
        spaceAfter=4,
    )

    # --------------------------------------------------------
    # DOCUMENT
    # --------------------------------------------------------

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        title=(
            f"{candidate_name or optimized.get('name', 'Candidate')}"
            " - Optimized Resume"
        ),
    )

    story = []

    # ========================================================
    # NAME
    # ========================================================

    name = (
        optimized.get("name")
        or candidate_name
        or "Candidate"
    )

    story.append(
        Paragraph(
            str(name),
            name_style,
        )
    )

    # ========================================================
    # CONTACT
    # ========================================================

    contact = (
        optimized.get("contact")
        or {}
    )

    contact_parts = []

    for key in [
        "email",
        "phone",
        "location",
        "linkedin",
        "portfolio",
    ]:

        value = contact.get(key)

        if value:
            contact_parts.append(
                str(value)
            )

    # Support old flat format too
    for key in [
        "email",
        "phone",
        "location",
        "linkedin",
        "portfolio",
    ]:

        if (
            not contact_parts
            and optimized.get(key)
        ):
            contact_parts.append(
                str(optimized.get(key))
            )

    if contact_parts:

        story.append(
            Paragraph(
                " | ".join(contact_parts),
                styles["Muted"],
            )
        )

    story.append(
        Spacer(1, 5)
    )

    story.append(
        HRFlowable(
            width="100%",
            color=colors.HexColor(
                "#E1E4EA"
            ),
            thickness=1,
        )
    )

    # ========================================================
    # SUMMARY
    # ========================================================

    summary = (
        optimized.get(
            "professional_summary"
        )
        or optimized.get("summary")
        or ""
    )

    if summary.strip():

        story.append(
            Paragraph(
                "PROFESSIONAL SUMMARY",
                heading_style,
            )
        )

        story.append(
            Paragraph(
                str(summary),
                styles["Body"],
            )
        )

    # ========================================================
    # SKILLS
    # ========================================================

    technical_skills = (
        optimized.get(
            "technical_skills"
        )
        or {}
    )

    # Support old format: skills = [{category, skills}]
    if not technical_skills:

        old_skills = (
            optimized.get("skills")
            or []
        )

        if isinstance(
            old_skills,
            list,
        ):

            converted = {}

            for item in old_skills:

                if not isinstance(
                    item,
                    dict,
                ):
                    continue

                category = (
                    item.get("category")
                    or "Skills"
                )

                values = (
                    item.get("skills")
                    or []
                )

                if values:
                    converted[
                        category
                    ] = values

            technical_skills = converted

    if technical_skills:

        story.append(
            Paragraph(
                "TECHNICAL SKILLS",
                heading_style,
            )
        )

        for category, items in (
            technical_skills.items()
        ):

            if not items:
                continue

            if not isinstance(
                items,
                list,
            ):
                items = [items]

            skill_text = ", ".join(
                str(item)
                for item in items
                if str(item).strip()
            )

            if not skill_text:
                continue

            story.append(
                Paragraph(
                    f"<b>{category}:</b> "
                    f"{skill_text}",
                    styles["Body"],
                )
            )

            story.append(
                Spacer(1, 2)
            )

    # ========================================================
    # EXPERIENCE
    # ========================================================

    experience = (
        optimized.get(
            "experience"
        )
        or []
    )

    if experience:

        story.append(
            Paragraph(
                "EXPERIENCE",
                heading_style,
            )
        )

        for entry in experience:

            if not isinstance(
                entry,
                dict,
            ):
                continue

            title = (
                entry.get("title")
                or ""
            )

            company = (
                entry.get("company")
                or ""
            )

            duration = (
                entry.get("duration")
                or entry.get("dates")
                or ""
            )

            header_parts = [
                str(x)
                for x in [
                    title,
                    company,
                ]
                if x
            ]

            if header_parts:

                story.append(
                    Paragraph(
                        " — ".join(
                            header_parts
                        ),
                        entry_title_style,
                    )
                )

            if duration:

                story.append(
                    Paragraph(
                        str(duration),
                        entry_meta_style,
                    )
                )

            bullets = (
                entry.get("bullets")
                or []
            )

            for bullet in bullets:

                if not str(
                    bullet
                ).strip():
                    continue

                story.append(
                    Paragraph(
                        f"•&nbsp;&nbsp;{bullet}",
                        styles["Body"],
                    )
                )

            story.append(
                Spacer(1, 4)
            )

    # ========================================================
    # PROJECTS
    # ========================================================

    projects = (
        optimized.get(
            "projects"
        )
        or []
    )

    if projects:

        story.append(
            Paragraph(
                "PROJECTS",
                heading_style,
            )
        )

        for project in projects:

            if not isinstance(
                project,
                dict,
            ):
                continue

            project_name = (
                project.get("name")
                or ""
            )

            if project_name:

                story.append(
                    Paragraph(
                        str(project_name),
                        entry_title_style,
                    )
                )

            duration = (
                project.get("duration")
                or ""
            )

            tech_stack = (
                project.get(
                    "tech_stack"
                )
                or project.get(
                    "technologies"
                )
                or []
            )

            if not isinstance(
                tech_stack,
                list,
            ):
                tech_stack = [
                    tech_stack
                ]

            meta_parts = []

            if duration:
                meta_parts.append(
                    str(duration)
                )

            if tech_stack:
                meta_parts.append(
                    ", ".join(
                        str(x)
                        for x in tech_stack
                        if str(x).strip()
                    )
                )

            if meta_parts:

                story.append(
                    Paragraph(
                        " | ".join(
                            meta_parts
                        ),
                        entry_meta_style,
                    )
                )

            bullets = (
                project.get(
                    "bullets"
                )
                or []
            )

            for bullet in bullets:

                if not str(
                    bullet
                ).strip():
                    continue

                story.append(
                    Paragraph(
                        f"•&nbsp;&nbsp;{bullet}",
                        styles["Body"],
                    )
                )

            story.append(
                Spacer(1, 4)
            )

    # ========================================================
    # EDUCATION
    # ========================================================

    education = (
        optimized.get(
            "education"
        )
        or []
    )

    if education:

        story.append(
            Paragraph(
                "EDUCATION",
                heading_style,
            )
        )

        for entry in education:

            if not isinstance(
                entry,
                dict,
            ):
                continue

            degree = (
                entry.get("degree")
                or ""
            )

            institution = (
                entry.get("institution")
                or ""
            )

            dates = (
                entry.get("dates")
                or entry.get("duration")
                or ""
            )

            line_parts = [
                str(x)
                for x in [
                    degree,
                    institution,
                ]
                if x
            ]

            if line_parts:

                story.append(
                    Paragraph(
                        " — ".join(
                            line_parts
                        ),
                        entry_title_style,
                    )
                )

            if dates:

                story.append(
                    Paragraph(
                        str(dates),
                        entry_meta_style,
                    )
                )

    # ========================================================
    # CERTIFICATIONS
    # ========================================================

    certifications = (
        optimized.get(
            "certifications"
        )
        or []
    )

    if certifications:

        story.append(
            Paragraph(
                "CERTIFICATIONS",
                heading_style,
            )
        )

        for cert in certifications:

            story.append(
                Paragraph(
                    f"•&nbsp;&nbsp;{cert}",
                    styles["Body"],
                )
            )

    # ========================================================
    # ACHIEVEMENTS
    # ========================================================

    achievements = (
        optimized.get(
            "achievements"
        )
        or []
    )

    if achievements:

        story.append(
            Paragraph(
                "ACHIEVEMENTS",
                heading_style,
            )
        )

        for achievement in achievements:

            story.append(
                Paragraph(
                    f"•&nbsp;&nbsp;{achievement}",
                    styles["Body"],
                )
            )

    # ========================================================
    # BUILD PDF
    # ========================================================

    doc.build(
        story,
        canvasmaker=functools.partial(
            _NumberedCanvas,
            footer_label=(
                f"{name} — Optimized Resume"
            ),
        ),
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# OPTIMIZED RESUME DOCX
# ============================================================

def generate_resume_docx(
    optimized: dict,
) -> bytes:

    """
    Generates DOCX from the SAME optimized dictionary.
    """

    doc = Document()

    for section in doc.sections:

        section.top_margin = Inches(
            0.65
        )

        section.bottom_margin = Inches(
            0.65
        )

        section.left_margin = Inches(
            0.7
        )

        section.right_margin = Inches(
            0.7
        )

    # --------------------------------------------------------
    # NAME
    # --------------------------------------------------------

    name = (
        optimized.get("name")
        or "Candidate"
    )

    p_name = doc.add_paragraph()

    p_name.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r_name = p_name.add_run(
        str(name)
    )

    r_name.font.size = Pt(18)
    r_name.font.bold = True

    # --------------------------------------------------------
    # CONTACT
    # --------------------------------------------------------

    contact = (
        optimized.get("contact")
        or {}
    )

    contact_parts = []

    for key in [
        "email",
        "phone",
        "location",
        "linkedin",
        "portfolio",
    ]:

        value = contact.get(key)

        if value:
            contact_parts.append(
                str(value)
            )

    if contact_parts:

        p_contact = (
            doc.add_paragraph()
        )

        p_contact.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p_contact.add_run(
            " | ".join(
                contact_parts
            )
        )

    # --------------------------------------------------------
    # SECTION HELPER
    # --------------------------------------------------------

    def add_section_title(title):

        p = doc.add_paragraph()

        p.paragraph_format.space_before = Pt(
            9
        )

        p.paragraph_format.space_after = Pt(
            2
        )

        r = p.add_run(
            title.upper()
        )

        r.bold = True

    # --------------------------------------------------------
    # SUMMARY
    # --------------------------------------------------------

    summary = (
        optimized.get(
            "professional_summary"
        )
        or optimized.get("summary")
        or ""
    )

    if summary:

        add_section_title(
            "Professional Summary"
        )

        doc.add_paragraph(
            str(summary)
        )

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    technical_skills = (
        optimized.get(
            "technical_skills"
        )
        or {}
    )

    if technical_skills:

        add_section_title(
            "Technical Skills"
        )

        for category, items in (
            technical_skills.items()
        ):

            if not items:
                continue

            if not isinstance(
                items,
                list,
            ):
                items = [items]

            p = doc.add_paragraph()

            r = p.add_run(
                f"{category}: "
            )

            r.bold = True

            p.add_run(
                ", ".join(
                    str(x)
                    for x in items
                    if str(x).strip()
                )
            )

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    experience = (
        optimized.get(
            "experience"
        )
        or []
    )

    if experience:

        add_section_title(
            "Experience"
        )

        for exp in experience:

            if not isinstance(
                exp,
                dict,
            ):
                continue

            title = (
                exp.get("title")
                or ""
            )

            company = (
                exp.get("company")
                or ""
            )

            duration = (
                exp.get("duration")
                or exp.get("dates")
                or ""
            )

            parts = [
                str(x)
                for x in [
                    title,
                    company,
                ]
                if x
            ]

            heading = " — ".join(
                parts
            )

            if duration:

                heading += (
                    f" ({duration})"
                )

            if heading:

                p = doc.add_paragraph()

                r = p.add_run(
                    heading
                )

                r.bold = True

            for bullet in (
                exp.get("bullets")
                or []
            ):

                doc.add_paragraph(
                    str(bullet),
                    style="List Bullet",
                )

    # --------------------------------------------------------
    # PROJECTS
    # --------------------------------------------------------

    projects = (
        optimized.get(
            "projects"
        )
        or []
    )

    if projects:

        add_section_title(
            "Projects"
        )

        for project in projects:

            if not isinstance(
                project,
                dict,
            ):
                continue

            project_name = (
                project.get("name")
                or ""
            )

            tech_stack = (
                project.get(
                    "tech_stack"
                )
                or project.get(
                    "technologies"
                )
                or []
            )

            if not isinstance(
                tech_stack,
                list,
            ):
                tech_stack = [
                    tech_stack
                ]

            heading = str(
                project_name
            )

            if tech_stack:

                heading += (
                    " | "
                    + ", ".join(
                        str(x)
                        for x in tech_stack
                        if str(x).strip()
                    )
                )

            if heading:

                p = doc.add_paragraph()

                r = p.add_run(
                    heading
                )

                r.bold = True

            for bullet in (
                project.get("bullets")
                or []
            ):

                doc.add_paragraph(
                    str(bullet),
                    style="List Bullet",
                )

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    education = (
        optimized.get(
            "education"
        )
        or []
    )

    if education:

        add_section_title(
            "Education"
        )

        for edu in education:

            if not isinstance(
                edu,
                dict,
            ):
                continue

            degree = (
                edu.get("degree")
                or ""
            )

            institution = (
                edu.get("institution")
                or ""
            )

            dates = (
                edu.get("dates")
                or edu.get("duration")
                or ""
            )

            parts = [
                str(x)
                for x in [
                    degree,
                    institution,
                ]
                if x
            ]

            line = " — ".join(
                parts
            )

            if dates:

                line += (
                    f" ({dates})"
                )

            doc.add_paragraph(
                line
            )

    # --------------------------------------------------------
    # CERTIFICATIONS
    # --------------------------------------------------------

    certifications = (
        optimized.get(
            "certifications"
        )
        or []
    )

    if certifications:

        add_section_title(
            "Certifications"
        )

        for cert in certifications:

            doc.add_paragraph(
                str(cert),
                style="List Bullet",
            )

    # --------------------------------------------------------
    # ACHIEVEMENTS
    # --------------------------------------------------------

    achievements = (
        optimized.get(
            "achievements"
        )
        or []
    )

    if achievements:

        add_section_title(
            "Achievements"
        )

        for achievement in achievements:

            doc.add_paragraph(
                str(achievement),
                style="List Bullet",
            )

    # --------------------------------------------------------
    # RETURN
    # --------------------------------------------------------

    buffer = io.BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()